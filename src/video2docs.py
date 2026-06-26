#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Video to Document Converter

This script converts YouTube videos or local video files to document formats (ODT, DOCX, PDF).
It extracts text from speech, identifies slides/images, and uses LLMs to organize the content.
"""

import os
import argparse
import tempfile
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Union
import logging
import re
import time
from dotenv import load_dotenv

# Video processing
import cv2
from pytubefix import YouTube
from pytubefix.cli import on_progress
from moviepy import VideoFileClip

# Audio processing
import speech_recognition as sr
from pydub import AudioSegment

# Image processing
from PIL import Image
import numpy as np
from skimage.metrics import structural_similarity as ssim

# Document generation
import docx
from docx.shared import Inches
from odf.opendocument import OpenDocumentText
from odf.text import P
from odf.style import Style, TextProperties
from odf.draw import Frame, Image as ODFImage
from fpdf import FPDF

# LLM and AI
import torch
from transformers import pipeline, AutoModelForSeq2SeqLM, AutoModelForCausalLM, AutoTokenizer
from langchain_community.llms import OpenAI
from langchain_openai import ChatOpenAI, AzureChatOpenAI
from langchain_community.llms.huggingface_pipeline import HuggingFacePipeline
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnablePassthrough

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()


class VideoProcessor:
    """Handles video downloading and processing."""

    def __init__(self, temp_dir: str = None):
        """Initialize the video processor.

        Args:
            temp_dir: Directory to store temporary files
        """
        self.temp_dir = temp_dir or tempfile.mkdtemp()
        logger.info(f"Using temporary directory: {self.temp_dir}")

    def download_youtube_video(self, url: str, max_retries: int = 3) -> str:
        """Download a YouTube video.

        Args:
            url: YouTube video URL
            max_retries: Maximum number of retry attempts

        Returns:
            Path to the downloaded video file
        """
        logger.info(f"Downloading YouTube video: {url}")

        # List of common YouTube domains to try if the original URL fails
        youtube_domains = [
            "www.youtube.com",
            "youtube.com",
            "youtu.be",
            "m.youtube.com",
        ]

        # Extract video ID from URL
        video_id = None
        if "youtube.com/watch" in url:
            # Format: https://www.youtube.com/watch?v=VIDEO_ID
            query_params = url.split("?")[1].split("&")
            for param in query_params:
                if param.startswith("v="):
                    video_id = param[2:]
                    break
        elif "youtu.be/" in url:
            # Format: https://youtu.be/VIDEO_ID
            video_id = url.split("youtu.be/")[1].split("?")[0]

        if not video_id:
            error_msg = f"Could not extract video ID from URL: {url}"
            logger.error(error_msg)
            raise ValueError(error_msg)

        # Try downloading with the original URL first
        urls_to_try = [url]

        # Add alternative URLs with different domains
        for domain in youtube_domains:
            if domain not in url:
                alt_url = f"https://{domain}/watch?v={video_id}"
                if alt_url != url:
                    urls_to_try.append(alt_url)

        last_exception = None
        for retry_count in range(max_retries):
            for try_url in urls_to_try:
                try:
                    logger.info(f"Attempt {retry_count + 1}/{max_retries} with URL: {try_url}")
                    yt = YouTube(try_url, on_progress_callback=on_progress)
                    logger.info(yt.title)
                    video = yt.streams.get_highest_resolution()

                    if not video:
                        logger.warning(f"No suitable video streams found for {try_url}")
                        continue

                    output_path = video.download(output_path=self.temp_dir)
                    logger.info(f"Downloaded video to: {output_path}")
                    return output_path
                except Exception as e:
                    logger.warning(f"Error downloading YouTube video from {try_url}: {e}")
                    last_exception = e
                    # Continue to the next URL or retry

            # Wait before retrying
            if retry_count < max_retries - 1:
                wait_time = 2 ** retry_count  # Exponential backoff
                logger.info(f"Waiting {wait_time} seconds before next retry...")
                time.sleep(wait_time)

        # If we get here, all attempts failed
        error_msg = f"Failed to download YouTube video after {max_retries} attempts with multiple URLs. Last error: {last_exception}"
        logger.error(error_msg)
        raise Exception(error_msg)

    def extract_audio(self, video_path: str) -> str:
        """Extract audio from video file.

        Args:
            video_path: Path to the video file

        Returns:
            Path to the extracted audio file
        """
        logger.info(f"Extracting audio from: {video_path}")
        try:
            video = VideoFileClip(video_path)
            audio_path = os.path.join(self.temp_dir, "audio.wav")
            video.audio.write_audiofile(audio_path, codec='pcm_s16le')
            logger.info(f"Extracted audio to: {audio_path}")
            return audio_path
        except Exception as e:
            logger.error(f"Error extracting audio: {e}")
            raise

    def extract_frames(self, video_path: str, interval: float = 1.0) -> List[Tuple[float, str]]:
        """Extract frames from video at specified intervals.

        Args:
            video_path: Path to the video file
            interval: Interval between frames in seconds

        Returns:
            List of tuples containing (timestamp, frame_path)
        """
        logger.info(f"Extracting frames from: {video_path}")
        frames = []
        try:
            video = cv2.VideoCapture(video_path)
            fps = video.get(cv2.CAP_PROP_FPS)
            frame_interval = int(fps * interval)

            success, frame = video.read()
            count = 0
            frame_count = 0

            while success:
                if count % frame_interval == 0:
                    timestamp = count / fps
                    frame_path = os.path.join(self.temp_dir, f"frame_{frame_count:04d}.jpg")
                    cv2.imwrite(frame_path, frame)
                    frames.append((timestamp, frame_path))
                    frame_count += 1

                success, frame = video.read()
                count += 1

            video.release()
            logger.info(f"Extracted {len(frames)} frames")
            return frames
        except Exception as e:
            logger.error(f"Error extracting frames: {e}")
            raise

    def detect_slides(self, frames: List[Tuple[float, str]], threshold: float = 0.92, min_gap_s: float = 10.0) -> List[Tuple[float, str]]:
        """Detect slides/images in the extracted frames.

        Args:
            frames: List of (timestamp, frame_path) tuples
            threshold: Similarity threshold — higher = fewer, more significant changes only
            min_gap_s: Minimum seconds between slides to ignore transient cursor/highlight changes

        Returns:
            List of (timestamp, frame_path) tuples for detected slides
        """
        logger.info("Detecting slides in frames")
        if not frames:
            return []

        slides = [frames[0]]

        for i in range(1, len(frames)):
            curr_ts = frames[i][0]
            if curr_ts - slides[-1][0] < min_gap_s:
                continue  # Too soon after last slide

            prev_img = cv2.imread(slides[-1][1], cv2.IMREAD_GRAYSCALE)
            curr_img = cv2.imread(frames[i][1], cv2.IMREAD_GRAYSCALE)

            if prev_img is None or curr_img is None:
                continue

            if prev_img.shape != curr_img.shape:
                curr_img = cv2.resize(curr_img, (prev_img.shape[1], prev_img.shape[0]))

            similarity = ssim(prev_img, curr_img)

            if similarity < threshold:
                slides.append(frames[i])
                logger.debug(f"Detected new slide at {curr_ts:.2f}s (similarity: {similarity:.2f})")

        logger.info(f"Detected {len(slides)} slides")
        return slides

    def extract_frame_at_time(self, video_path: str, timestamp: float, filename: str):
        """Extract a single frame at the given timestamp (seconds) from a video."""
        try:
            video = cv2.VideoCapture(video_path)
            fps = video.get(cv2.CAP_PROP_FPS) or 25.0
            video.set(cv2.CAP_PROP_POS_FRAMES, int(timestamp * fps))
            success, frame = video.read()
            video.release()
            if success:
                cv2.imwrite(filename, frame)
                return filename
            return None
        except Exception as e:
            logger.error(f"Error extracting frame at {timestamp}s: {e}")
            return None


class AudioProcessor:
    """Handles audio transcription."""

    def __init__(self, use_gpu: bool = True):
        """Initialize the audio processor.

        Args:
            use_gpu: Whether to use GPU for processing
        """
        self.use_gpu = use_gpu and torch.cuda.is_available()
        self.recognizer = sr.Recognizer()
        logger.info(f"Initialized audio processor (GPU: {self.use_gpu})")

    def transcribe_audio(self, audio_path: str, chunk_size: int = 600000, language: str = "en-US") -> List[Dict[str, Union[str, float, float]]]:
        """Transcribe audio file to text using Azure Whisper, local Whisper, or Google fallback."""
        logger.info(f"Transcribing audio: {audio_path}")
        use_whisper = (
            (os.getenv("AZURE_WHISPER_ENDPOINT") and os.getenv("AZURE_WHISPER_KEY"))
            or (os.getenv("OPENAI_API_KEY") and os.getenv("OPENAI_API_BASE"))
        )
        if use_whisper:
            return self._transcribe_whisper(audio_path, chunk_size, language, None, None)
        logger.warning("No Whisper config found — falling back to Google SR (poor quality)")
        return self._transcribe_google(audio_path, chunk_size, language)

    def _transcribe_whisper(self, audio_path: str, chunk_size: int, language: str, api_key: str, api_base: str) -> List[Dict]:
        """Transcribe using Azure Whisper API if configured, otherwise local model."""
        azure_whisper_endpoint = os.getenv("AZURE_WHISPER_ENDPOINT")
        azure_whisper_key = os.getenv("AZURE_WHISPER_KEY")
        if azure_whisper_endpoint and azure_whisper_key:
            return self._transcribe_azure_whisper(audio_path, chunk_size, language, azure_whisper_endpoint, azure_whisper_key)
        return self._transcribe_local_whisper(audio_path, language)

    def _transcribe_azure_whisper(self, audio_path: str, chunk_size: int, language: str, endpoint: str, api_key: str) -> List[Dict]:
        """Transcribe using Azure-hosted Whisper REST API — 25MB chunks, no local GPU needed."""
        import openai as _openai
        lang_code = language.split("-")[0]
        client = _openai.AzureOpenAI(
            api_key=api_key,
            azure_endpoint=endpoint,
            api_version="2024-06-01",
        )
        audio = AudioSegment.from_file(audio_path)
        duration = len(audio)
        chunks = []

        for start_ms in range(0, duration, chunk_size):
            end_ms = min(start_ms + chunk_size, duration)
            chunk = audio[start_ms:end_ms]
            chunk_path = f"{audio_path}_chunk_{start_ms}.mp3"
            try:
                chunk.export(chunk_path, format="mp3", bitrate="64k")
                file_size = os.path.getsize(chunk_path)
                if file_size > 24 * 1024 * 1024:
                    logger.warning(f"Chunk {start_ms}-{end_ms}ms too large ({file_size//1024}KB), skipping")
                    chunks.append({"text": "", "start_time": start_ms / 1000.0, "end_time": end_ms / 1000.0})
                    continue
                logger.info(f"Azure Whisper transcribing {start_ms//1000}s-{end_ms//1000}s ({file_size//1024}KB)")
                with open(chunk_path, "rb") as f:
                    result = client.audio.transcriptions.create(
                        model="whisper",
                        file=f,
                        language=lang_code,
                        response_format="verbose_json",
                        timestamp_granularities=["segment"],
                    )
                chunk_offset = start_ms / 1000.0
                if hasattr(result, 'segments') and result.segments:
                    for seg in result.segments:
                        chunks.append({
                            "text": seg.text.strip(),
                            "start_time": chunk_offset + seg.start,
                            "end_time": chunk_offset + seg.end,
                        })
                    logger.info(f"Azure Whisper chunk: {len(result.segments)} segments, first: {result.segments[0].text[:60]}")
                else:
                    # Fallback: single chunk with coarse timestamp
                    chunks.append({"text": result.text.strip(), "start_time": chunk_offset, "end_time": end_ms / 1000.0})
                    logger.info(f"Azure Whisper chunk (no segments): {result.text[:80]}")
            except Exception as e:
                logger.error(f"Azure Whisper error {start_ms}-{end_ms}ms: {e}")
                chunks.append({"text": "", "start_time": start_ms / 1000.0, "end_time": end_ms / 1000.0})
            finally:
                if os.path.exists(chunk_path):
                    try:
                        os.remove(chunk_path)
                    except Exception:
                        pass

        logger.info(f"Azure Whisper transcribed {len(chunks)} chunks")
        return chunks

    def _transcribe_local_whisper(self, audio_path: str, language: str) -> List[Dict]:
        """Transcribe using local Whisper via transformers pipeline."""
        from transformers import pipeline as hf_pipeline
        lang_code = language.split("-")[0]
        whisper_model = os.getenv("WHISPER_MODEL", "openai/whisper-large-v3")
        logger.info(f"Loading local Whisper model: {whisper_model}")
        asr = hf_pipeline(
            "automatic-speech-recognition",
            model=whisper_model,
            generate_kwargs={"language": lang_code, "task": "transcribe"},
            chunk_length_s=30,
            stride_length_s=5,
        )
        logger.info(f"Running local Whisper on {audio_path}")
        result = asr(audio_path, return_timestamps=True)
        chunks_raw = result.get("chunks", [])
        if chunks_raw:
            chunks = [
                {"text": c["text"].strip(), "start_time": c["timestamp"][0] or 0.0, "end_time": c["timestamp"][1] or 0.0}
                for c in chunks_raw
            ]
        else:
            full_text = result.get("text", "")
            audio = AudioSegment.from_file(audio_path)
            chunks = [{"text": full_text, "start_time": 0.0, "end_time": len(audio) / 1000.0}]
        logger.info(f"Local Whisper transcribed {len(chunks)} segments")
        return chunks

    def _transcribe_google(self, audio_path: str, chunk_size: int, language: str) -> List[Dict]:
        """Fallback: Google free speech recognition."""
        logger.info("Using Google speech recognition (fallback)")
        audio = AudioSegment.from_file(audio_path)
        duration = len(audio)
        chunks = []

        for start_time in range(0, duration, min(chunk_size, 60000)):
            end_time = min(start_time + 60000, duration)
            chunk = audio[start_time:end_time]
            chunk_path = f"{audio_path}_chunk_{start_time}_{end_time}.wav"
            try:
                chunk.export(chunk_path, format="wav")
                with sr.AudioFile(chunk_path) as source:
                    audio_data = self.recognizer.record(source)
                    try:
                        text = self.recognizer.recognize_google(audio_data, language=language)
                    except sr.UnknownValueError:
                        text = ""
                    except sr.RequestError as e:
                        logger.error(f"Google SR error: {e}")
                        text = ""
                chunks.append({"text": text, "start_time": start_time / 1000.0, "end_time": end_time / 1000.0})
            finally:
                if os.path.exists(chunk_path):
                    try:
                        os.remove(chunk_path)
                    except Exception:
                        pass

        logger.info(f"Google transcribed {len(chunks)} chunks")
        return chunks


class LLMProcessor:
    """Handles LLM processing for content organization."""

    def __init__(self, model_name: str = None, use_openai: bool = False, use_gpu: bool = True):
        """Initialize the LLM processor.

        Args:
            model_name: Name of the HuggingFace model to use
            use_openai: Whether to use OpenAI API
            use_gpu: Whether to use GPU for processing
        """
        self.use_openai = use_openai
        self.use_gpu = use_gpu and torch.cuda.is_available()

        if use_openai and os.getenv("OPENAI_API_KEY"):
            chat_model = model_name or os.getenv("VIDEO2DOCS_LLM_MODEL", "gpt-4.1")
            azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
            logger.info(f"Using {'Azure' if azure_endpoint else 'OpenAI-compatible'} API: {chat_model}")
            if azure_endpoint:
                self.llm = AzureChatOpenAI(
                    azure_deployment=chat_model,
                    azure_endpoint=azure_endpoint,
                    api_key=os.getenv("OPENAI_API_KEY"),
                    api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01"),
                    request_timeout=120,
                    max_retries=3,
                    streaming=True,
                )
            else:
                self.llm = ChatOpenAI(
                    model_name=chat_model,
                    openai_api_key=os.getenv("OPENAI_API_KEY"),
                    openai_api_base=os.getenv("OPENAI_API_BASE"),
                    request_timeout=120,
                    max_retries=3,
                    streaming=True,
                )
        else:
            # Default to a HuggingFace model (allow override via .env)
            env_model = os.getenv("VIDEO2DOCS_LLM_MODEL", "").strip()
            self.model_name = (model_name or env_model or "google/flan-t5-large")
            if not model_name and env_model:
                logger.info(f"Using HuggingFace model from env VIDEO2DOCS_LLM_MODEL: {self.model_name}")
            else:
                logger.info(f"Using HuggingFace model: {self.model_name}")

            # Set device based on GPU availability and preference
            device = "cuda" if self.use_gpu else "cpu"
            logger.info(f"Using device: {device}")

            # Download and load the model locally (try Seq2Seq first, then Causal LM)
            logger.info(f"Downloading and loading model: {self.model_name}")
            model = None
            tokenizer = None
            is_seq2seq = False
            try:
                model = AutoModelForSeq2SeqLM.from_pretrained(self.model_name)
                is_seq2seq = True
            except Exception as e1:
                logger.debug(f"Seq2Seq load failed for {self.model_name}: {e1}")
                try:
                    model = AutoModelForCausalLM.from_pretrained(self.model_name)
                    is_seq2seq = False
                except Exception as e2:
                    logger.error(f"Error loading model {self.model_name} as Seq2Seq or Causal: {e1} | {e2}")
                    logger.error("Please check your internet connection and ensure you have enough disk space.")
                    logger.error("If the error persists, try a different model or check if it is available on Hugging Face Hub.")
                    raise RuntimeError(f"Failed to load model {self.model_name}: {e2}") from e2

            try:
                tokenizer = AutoTokenizer.from_pretrained(self.model_name)
            except Exception as e_tok:
                logger.error(f"Error loading tokenizer for {self.model_name}: {e_tok}")
                raise RuntimeError(f"Failed to load tokenizer for model {self.model_name}: {e_tok}") from e_tok

            # Move model to GPU if available
            model = model.to(device)
            logger.info(f"Successfully loaded model {self.model_name} to {device} (seq2seq={is_seq2seq})")

            # Create a text generation pipeline (task depends on architecture)
            task = "text2text-generation" if is_seq2seq else "text-generation"
            text_generation_pipeline = pipeline(
                task,
                model=model,
                tokenizer=tokenizer,
                device=0 if device == "cuda" else -1,
                max_length=512,
                temperature=0.1,
                num_return_sequences=1,  # Only return one sequence
                do_sample=True  # Enable sampling for more creative outputs
            )

            # Create a LangChain pipeline
            self.llm = HuggingFacePipeline(pipeline=text_generation_pipeline)

    def organize_content(self, transcription: List[Dict], slides: List[Tuple[float, str]]) -> Dict:
        """Organize content using LLM.

        Args:
            transcription: List of transcription chunks
            slides: List of detected slides

        Returns:
            Organized document structure
        """
        logger.info("Organizing content with LLM")

        # Build timestamped transcript
        transcript_lines = []
        for chunk in transcription:
            if chunk.get("text", "").strip():
                t = int(chunk["start_time"])
                mins, secs = divmod(t, 60)
                transcript_lines.append(f"[{mins:02d}:{secs:02d}] {chunk['text'].strip()}")
        timestamped_transcript = "\n".join(transcript_lines)

        prompt = PromptTemplate(
            input_variables=["transcript"],
            template="""
            You are an AI assistant that organizes video content into a structured document.

            TIMESTAMPED TRANSCRIPT (format [MM:SS] text):
            {transcript}

            Instructions:
            - Divide the transcript into logical sections based on topic changes
            - For each section record the start_time and end_time in total seconds (integers) from the [MM:SS] timestamps
            - bullet_points are the PRIMARY content — detailed and informative, 4-6 points per section
            - content is a single brief intro sentence for the section
            - Preserve chronological order — sections must follow the video timeline

            Respond ONLY with JSON, no markdown fences:
            {{
                "title": "Document Title",
                "summary": "Executive summary 2-3 sentences",
                "sections": [
                    {{
                        "heading": "Section Heading",
                        "start_time": 0,
                        "end_time": 120,
                        "content": "One sentence intro.",
                        "bullet_points": ["Detailed point 1", "Detailed point 2"]
                    }}
                ]
            }}
            """
        )

        text_limit = 20000 if self.use_openai else 1500
        formatted_prompt = prompt.format(transcript=timestamped_transcript[:text_limit])

        # Invoke the LLM
        raw_result = self.llm.invoke(formatted_prompt)

        if hasattr(raw_result, 'content'):
            result = raw_result.content
        elif isinstance(raw_result, list) and len(raw_result) > 0 and 'generated_text' in raw_result[0]:
            result = raw_result[0]['generated_text']
        else:
            result = str(raw_result)

        logger.debug(f"Processed result: {result[:100]}...")

        # Parse the result
        try:
            import json
            # Strip markdown fences if present
            json_match = re.search(r'```(?:json)?\s*(.*?)\s*```', result, re.DOTALL)
            result = json_match.group(1) if json_match else result.strip()
            document_structure = json.loads(result)

            # Inject slides by timestamp — deterministic, no LLM guessing
            slide_times = [(ts, path) for ts, path in slides]
            used = set()
            for section in document_structure.get("sections", []):
                sec_start = float(section.get("start_time", 0))
                sec_end = float(section.get("end_time", 999999))
                section_slides = []
                for i, (ts, path) in enumerate(slide_times):
                    if i not in used and sec_start <= ts < sec_end:
                        section_slides.append(i + 1)  # 1-based for document generator
                        used.add(i)
                section["slide_indices"] = section_slides

            # Any slides outside all section ranges → assign to nearest section
            for i, (ts, _) in enumerate(slide_times):
                if i not in used and document_structure.get("sections"):
                    nearest = min(
                        range(len(document_structure["sections"])),
                        key=lambda j: abs(float(document_structure["sections"][j].get("start_time", 0)) - ts)
                    )
                    document_structure["sections"][nearest].setdefault("slide_indices", []).append(i + 1)
                    used.add(i)

            logger.info("Successfully organised content with timestamp-matched slides")
            return document_structure
        except Exception as e:
            logger.error(f"Error parsing LLM response: {e}")
            logger.debug(f"Raw LLM response: {result}")
            full_text = " ".join([c["text"] for c in transcription])
            return {
                "title": "Transcribed Video",
                "summary": full_text[:500] + "...",
                "sections": [{"heading": "Full Transcription", "start_time": 0, "end_time": 999999,
                               "content": full_text, "bullet_points": [], "slide_indices": list(range(1, len(slides)+1))}]
            }


class DocumentGenerator:
    """Generates documents in various formats."""

    def __init__(self, output_dir: str = "."):
        """Initialize the document generator.

        Args:
            output_dir: Directory to save output documents
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        logger.info(f"Document generator initialized (output dir: {output_dir})")

    def generate_docx(self, content: Dict, slides: List[Tuple[float, str]], output_path: str) -> str:
        """Generate a DOCX document.

        Args:
            content: Organized document content
            slides: List of detected slides
            output_path: Path to save the document

        Returns:
            Path to the generated document
        """
        logger.info(f"Generating DOCX document: {output_path}")
        doc = docx.Document()

        # Add title
        doc.add_heading(content["title"], level=0)

        # Add summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(content["summary"])

        # Add sections — text leads, screenshots follow at end of section
        for section in content["sections"]:
            doc.add_heading(section["heading"], level=1)

            if section.get("content", "").strip():
                doc.add_paragraph(section["content"])

            if section.get("bullet_points"):
                for point in section["bullet_points"]:
                    doc.add_paragraph(point, style='List Bullet')

            # Screenshots after the text content they illustrate
            for slide_num in section.get("slide_indices", []):
                slide_index = slide_num - 1
                if 0 <= slide_index < len(slides):
                    try:
                        doc.add_picture(slides[slide_index][1], width=Inches(6))
                        doc.add_paragraph()
                    except Exception as e:
                        logger.error(f"Error adding slide {slide_index}: {e}")

        # Save document
        doc.save(output_path)
        logger.info(f"DOCX document saved to: {output_path}")
        return output_path

    def generate_odt(self, content: Dict, slides: List[Tuple[float, str]], output_path: str) -> str:
        """Generate an ODT document.

        Args:
            content: Organized document content
            slides: List of detected slides
            output_path: Path to save the document

        Returns:
            Path to the generated document
        """
        logger.info(f"Generating ODT document: {output_path}")
        doc = OpenDocumentText()

        # Add styles
        heading_style = Style(name="Heading", family="paragraph")
        heading_style.addElement(TextProperties(attributes={'fontsize': "16pt", 'fontweight': "bold"}))
        doc.styles.addElement(heading_style)

        # Add title
        title = P(stylename=heading_style)
        title.addText(content["title"])
        doc.text.addElement(title)

        # Add summary
        summary_heading = P(stylename=heading_style)
        summary_heading.addText("Executive Summary")
        doc.text.addElement(summary_heading)

        summary = P()
        summary.addText(content["summary"])
        doc.text.addElement(summary)

        # Add a section for slides/images if available
        if slides:
            slides_heading = P(stylename=heading_style)
            slides_heading.addText("Slides/Images")
            doc.text.addElement(slides_heading)

            for i, (timestamp, slide_path) in enumerate(slides):
                try:
                    # Add a caption for the slide
                    caption = P()
                    caption.addText(f"Slide {i+1} (Timestamp: {timestamp:.2f}s)")
                    doc.text.addElement(caption)

                    # Add image
                    frame = Frame(width="6in", height="4in")
                    img = ODFImage(href=slide_path)
                    frame.addElement(img)
                    doc.text.addElement(frame)

                    # Add some space after each image
                    spacer = P()
                    doc.text.addElement(spacer)
                except Exception as e:
                    logger.error(f"Error adding slide {i} from {slide_path}: {e}")

        # Add sections
        for section in content["sections"]:
            section_heading = P(stylename=heading_style)
            section_heading.addText(section["heading"])
            doc.text.addElement(section_heading)

            # Process content with slide markers
            content_parts = re.split(r'(\[SLIDE \d+\])', section["content"])
            for part in content_parts:
                slide_match = re.match(r'\[SLIDE (\d+)\]', part)
                if slide_match:
                    slide_index = int(slide_match.group(1))
                    if 0 <= slide_index < len(slides):
                        try:
                            # Add image
                            frame = Frame(width="6in", height="4in")
                            img = ODFImage(href=slides[slide_index][1])
                            frame.addElement(img)
                            doc.text.addElement(frame)
                        except Exception as e:
                            logger.error(f"Error adding slide {slide_index} from {slides[slide_index][1]}: {e}")
                else:
                    if part.strip():
                        p = P()
                        p.addText(part)
                        doc.text.addElement(p)

            # Add bullet points
            if section.get("bullet_points"):
                for point in section["bullet_points"]:
                    bullet = P(stylename="List")
                    bullet.addText("• " + point)
                    doc.text.addElement(bullet)

        # Save document
        doc.save(output_path)
        logger.info(f"ODT document saved to: {output_path}")
        return output_path

    def generate_pdf(self, content: Dict, slides: List[Tuple[float, str]], output_path: str) -> str:
        """Generate a PDF document.

        Args:
            content: Organized document content
            slides: List of detected slides
            output_path: Path to save the document

        Returns:
            Path to the generated document
        """
        logger.info(f"Generating PDF document: {output_path}")
        pdf = FPDF()
        pdf.add_page()

        # Set up fonts
        pdf.set_font("Arial", "B", 16)

        # Add title
        pdf.cell(0, 10, content["title"], ln=True, align="C")
        pdf.ln(10)

        # Add summary
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Executive Summary", ln=True)
        pdf.set_font("Arial", "", 12)
        pdf.multi_cell(0, 10, content["summary"])
        pdf.ln(10)

        # Add a section for slides/images if available
        if slides:
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "Slides/Images", ln=True)
            pdf.set_font("Arial", "", 12)

            for i, (timestamp, slide_path) in enumerate(slides):
                try:
                    # Add a caption for the slide
                    pdf.multi_cell(0, 10, f"Slide {i+1} (Timestamp: {timestamp:.2f}s)")

                    # Add image
                    pdf.image(slide_path, x=10, w=190)

                    # Add some space after each image
                    pdf.ln(5)
                except Exception as e:
                    logger.error(f"Error adding slide {i} from {slide_path}: {e}")

            pdf.ln(10)

        # Add sections
        for section in content["sections"]:
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, section["heading"], ln=True)
            pdf.set_font("Arial", "", 12)

            # Process content with slide markers
            content_parts = re.split(r'(\[SLIDE \d+\])', section["content"])
            for part in content_parts:
                slide_match = re.match(r'\[SLIDE (\d+)\]', part)
                if slide_match:
                    slide_index = int(slide_match.group(1))
                    if 0 <= slide_index < len(slides):
                        try:
                            # Add image
                            pdf.image(slides[slide_index][1], x=10, w=190)
                        except Exception as e:
                            logger.error(f"Error adding slide {slide_index} from {slides[slide_index][1]}: {e}")
                else:
                    if part.strip():
                        pdf.multi_cell(0, 10, part)

            # Add bullet points
            if section.get("bullet_points"):
                for point in section["bullet_points"]:
                    pdf.cell(10, 10, "•", ln=0)
                    pdf.multi_cell(0, 10, point)

        # Save document
        pdf.output(output_path)
        logger.info(f"PDF document saved to: {output_path}")
        return output_path


class Video2Docs:
    """Main class for video to document conversion."""

    def __init__(self, output_dir: str = "output", temp_dir: str = None, use_gpu: bool = True, model_name: str = None):
        """Initialize the converter.

        Args:
            output_dir: Directory to save output documents
            temp_dir: Directory for temporary files
            use_gpu: Whether to use GPU for processing
            model_name: Optional Hugging Face model name to use for summarization
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

        self.temp_dir = temp_dir or os.path.join(output_dir, "temp")
        os.makedirs(self.temp_dir, exist_ok=True)

        self.use_gpu = use_gpu and torch.cuda.is_available()
        if self.use_gpu:
            logger.info("GPU is available and will be used")
        else:
            logger.info("GPU is not available, using CPU")

        # Initialize components
        self.video_processor = VideoProcessor(temp_dir=self.temp_dir)
        self.audio_processor = AudioProcessor(use_gpu=self.use_gpu)
        use_openai = bool(os.getenv("OPENAI_API_KEY") and os.getenv("USE_OPENAI", "1") != "0")
        self.llm_processor = LLMProcessor(model_name=model_name, use_openai=use_openai, use_gpu=self.use_gpu)
        self.document_generator = DocumentGenerator(output_dir=self.output_dir)

    def process(self, input_path: str, output_format: str = "docx", output_name: Optional[str] = None, language: Optional[str] = None, progress_callback=None, cancel_event=None) -> str:
        """Process a video and convert it to a document.

        Args:
            input_path: Path or URL to the video
            output_format: Output document format (docx, odt, pdf)
            output_name: Optional base name (without extension) for the output file. If not provided,
                the name is derived from the input video file or video title.

        Returns:
            Path to the generated document
        """
        start_time = time.time()
        logger.info(f"Starting conversion of: {input_path}")

        # Progress/ETA helpers
        weights = {
            "download": 0.10,
            "extract_audio": 0.05,
            "detect_slides": 0.10,
            "transcribe": 0.55,
            "organize_content": 0.05,
            "generate_document": 0.15,
        }
        base = 0.0

        def report(step: str, step_progress: float):
            if progress_callback:
                total_progress = (base + weights.get(step, 0.0) * float(step_progress or 0.0)) * 100.0
                elapsed = time.time() - start_time
                eta = int(elapsed * (100.0 - total_progress) / max(total_progress, 1e-6)) if total_progress > 0 else None
                try:
                    progress_callback(step, round(total_progress, 1), eta)
                except Exception:
                    pass

        def check_cancel():
            if cancel_event is not None:
                try:
                    if cancel_event.is_set():
                        from .jobs import CancelledError
                        raise CancelledError()
                except AttributeError:
                    # Non-standard event
                    pass

        try:
            # Step 1: Get the video file
            report("download", 0.0)
            if input_path.startswith(("http://", "https://")) and "youtube" in input_path:
                try:
                    video_path = self.video_processor.download_youtube_video(input_path)
                    base_name = os.path.basename(video_path).split(".")[0]
                except Exception as youtube_error:
                    # Provide a more helpful error message for YouTube download failures
                    error_msg = str(youtube_error)
                    if "HTTP Error 400: Bad Request" in error_msg:
                        raise ValueError(
                            f"Failed to download YouTube video: {input_path}\n"
                            f"This could be due to one of the following reasons:\n"
                            f"1. The video might be unavailable, private, or age-restricted\n"
                            f"2. YouTube may have changed its API, affecting the pytube library\n"
                            f"3. There might be network restrictions or proxy issues\n\n"
                            f"Possible solutions:\n"
                            f"- Try downloading a local video file instead of a YouTube URL\n"
                            f"- Check if pytube has been updated to address YouTube API changes\n"
                            f"- Try running the script on a different network\n"
                            f"- Consider using youtube-dl or yt-dlp as alternative download methods\n\n"
                            f"Original error: {error_msg}"
                        ) from youtube_error
                    else:
                        # For other YouTube-related errors
                        raise ValueError(
                            f"Error downloading YouTube video: {input_path}\n"
                            f"Please check your internet connection and try again.\n"
                            f"Original error: {error_msg}"
                        ) from youtube_error
            else:
                # Local video file
                if not os.path.exists(input_path):
                    raise FileNotFoundError(f"Video file not found: {input_path}")

                video_path = input_path
                base_name = os.path.basename(video_path).split(".")[0]

            # Download complete
            base += weights.get("download", 0.0)
            report("download", 1.0)
            check_cancel()

            # Override base name if a custom output name is provided
            if output_name:
                base_name = os.path.splitext(output_name)[0]

            # Cache key: sha256 of first 4MB of video (fast, stable)
            import hashlib, json as _json
            def _video_hash(path):
                h = hashlib.sha256()
                with open(path, "rb") as f:
                    h.update(f.read(4 * 1024 * 1024))
                return h.hexdigest()[:16]

            cache_dir = os.path.join(self.output_dir, ".cache")
            os.makedirs(cache_dir, exist_ok=True)
            vid_hash = _video_hash(video_path)
            transcription_cache = os.path.join(cache_dir, f"{vid_hash}_transcription.json")
            slides_cache = os.path.join(cache_dir, f"{vid_hash}_slides.json")

            # Step 2: Extract audio
            report("extract_audio", 0.0)
            audio_path = self.video_processor.extract_audio(video_path)
            base += weights.get("extract_audio", 0.0)
            report("extract_audio", 1.0)
            check_cancel()

            # Step 3: Transcribe audio (cached)
            if os.path.exists(transcription_cache):
                logger.info(f"Loading transcription from cache: {transcription_cache}")
                with open(transcription_cache) as f:
                    transcription = _json.load(f)
                base += weights.get("transcribe", 0.0)
            else:
                report("transcribe", 0.0)
                transcription = self.audio_processor.transcribe_audio(audio_path, language=language) if language else self.audio_processor.transcribe_audio(audio_path)
                base += weights.get("transcribe", 0.0)
                report("transcribe", 1.0)
                with open(transcription_cache, "w") as f:
                    _json.dump(transcription, f)
                logger.info(f"Transcription cached: {transcription_cache}")
            check_cancel()

            # Step 4: Organize content with LLM (no slides yet — sections drive screenshot timing)
            report("organize_content", 0.0)
            content = self.llm_processor.organize_content(transcription, [])
            base += weights.get("organize_content", 0.0)
            report("organize_content", 1.0)
            check_cancel()

            # Step 5: Extract one screenshot per section at the section midpoint (cached)
            if os.path.exists(slides_cache):
                logger.info(f"Loading slides from cache: {slides_cache}")
                with open(slides_cache) as f:
                    slides_data = _json.load(f)
                slides = [(s["timestamp"], s["path"]) for s in slides_data if os.path.exists(s["path"])]
                if not slides:
                    logger.info("Cached slide images missing, re-extracting")
                    os.remove(slides_cache)
            if not os.path.exists(slides_cache):
                report("detect_slides", 0.0)
                slides = []
                # Get actual video duration to clamp timestamps
                _vid_cap = cv2.VideoCapture(video_path)
                _vid_fps = _vid_cap.get(cv2.CAP_PROP_FPS) or 25.0
                _vid_frames = _vid_cap.get(cv2.CAP_PROP_FRAME_COUNT)
                _vid_cap.release()
                video_duration = (_vid_frames / _vid_fps) - 1.0 if _vid_frames > 0 else 999999.0
                for i, section in enumerate(content.get("sections", [])):
                    sec_start = float(section.get("start_time", 0))
                    sec_end = float(section.get("end_time", sec_start + 60))
                    midpoint = min((sec_start + sec_end) / 2.0, video_duration)
                    frame_path = os.path.join(self.video_processor.temp_dir, f"section_{i:03d}.jpg")
                    result_frame = self.video_processor.extract_frame_at_time(video_path, midpoint, frame_path)
                    if result_frame:
                        slides.append((midpoint, result_frame))
                    else:
                        logger.warning(f"Could not extract frame for section {i} at {midpoint:.1f}s")
                base += weights.get("detect_slides", 0.0)
                report("detect_slides", 1.0)
                check_cancel()
                with open(slides_cache, "w") as f:
                    _json.dump([{"timestamp": ts, "path": p} for ts, p in slides], f)
                logger.info(f"Section screenshots cached: {slides_cache} ({len(slides)} frames)")
            else:
                base += weights.get("detect_slides", 0.0)

            # Assign each section its corresponding screenshot (1-based index)
            for i, section in enumerate(content.get("sections", [])):
                section["slide_indices"] = [i + 1] if i < len(slides) else []

            # Step 6: Generate document
            output_path = os.path.join(self.output_dir, f"{base_name}.{output_format}")

            report("generate_document", 0.0)
            if output_format.lower() == "docx":
                result_path = self.document_generator.generate_docx(content, slides, output_path)
            elif output_format.lower() == "odt":
                result_path = self.document_generator.generate_odt(content, slides, output_path)
            elif output_format.lower() == "pdf":
                result_path = self.document_generator.generate_pdf(content, slides, output_path)
            else:
                raise ValueError(f"Unsupported output format: {output_format}")

            base += weights.get("generate_document", 0.0)
            report("generate_document", 1.0)
            check_cancel()

            elapsed_time = time.time() - start_time
            logger.info(f"Conversion completed in {elapsed_time:.2f} seconds")
            logger.info(f"Output document: {result_path}")

            return result_path

        except ValueError as ve:
            # Re-raise ValueError exceptions (including our custom YouTube errors)
            logger.error(f"Error during conversion: {ve}")
            raise
        except FileNotFoundError as fnf:
            # Re-raise FileNotFoundError exceptions
            logger.error(f"Error during conversion: {fnf}")
            raise
        except Exception as e:
            # Generic error handling for other exceptions
            logger.error(f"Error during conversion: {e}")
            raise


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(description="Convert video to document")
    parser.add_argument("input", help="YouTube URL or path to local video file")
    parser.add_argument(
        "--format", "-f", 
        choices=["docx", "odt", "pdf"], 
        default="docx",
        help="Output document format"
    )
    parser.add_argument(
        "--output-dir", "-o",
        default="output",
        help="Directory to save output documents"
    )
    parser.add_argument(
        "--temp-dir", "-t",
        default=None,
        help="Directory for temporary files"
    )
    parser.add_argument(
        "--no-gpu",
        action="store_true",
        help="Disable GPU usage"
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Enable verbose logging"
    )
    parser.add_argument(
        "--language", "-l",
        default=os.environ.get("VIDEO2DOCS_LANGUAGE", "en-US"),
        help="Language code (BCP-47) for speech recognition, e.g., en-US, ru-RU"
    )

    args = parser.parse_args()

    # Configure logging
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    # Create converter
    converter = Video2Docs(
        output_dir=args.output_dir,
        temp_dir=args.temp_dir,
        use_gpu=not args.no_gpu
    )

    # Process video
    try:
        output_path = converter.process(args.input, args.format, language=args.language)
        print(f"Document generated: {output_path}")
        return 0
    except Exception as e:
        logger.error(f"Error: {e}")
        return 1


if __name__ == "__main__":
    exit(main())
